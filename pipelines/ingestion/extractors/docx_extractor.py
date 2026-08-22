from pathlib import Path

from docx import Document

from pipelines.ingestion.models import ExtractedPage


class DOCXExtractor:
    """Extract paragraphs, headings, headers, and table cells from DOCX files."""

    async def extract(self, file_path: str | Path) -> list[ExtractedPage]:
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"DOCX file not found: {path}")

        document = Document(str(path))

        pages: list[ExtractedPage] = []
        page_number = 1

        # Extract document paragraphs and preserve heading hierarchy.
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()

            if not text:
                continue

            style_name = paragraph.style.name if paragraph.style else ""

            metadata = {
                "source": "paragraph",
            }

            if style_name.startswith("Heading"):
                level = self._heading_level(style_name)

                metadata.update(
                    {
                        "level": f"h{level}",
                        "section": text,
                    }
                )

            pages.append(
                ExtractedPage(
                    page_number=page_number,
                    content=text,
                    content_type="text",
                    metadata=metadata,
                )
            )

            page_number += 1

        # Extract headers separately.
        for section_index, section in enumerate(
            document.sections,
            start=1,
        ):
            header_text = "\n".join(
                paragraph.text.strip()
                for paragraph in section.header.paragraphs
                if paragraph.text.strip()
            )

            if header_text:
                pages.append(
                    ExtractedPage(
                        page_number=page_number,
                        content=header_text,
                        content_type="text",
                        metadata={
                            "source": "header",
                            "section_index": section_index,
                        },
                    )
                )

                page_number += 1

        # Extract table cells separately.
        for table_index, table in enumerate(
            document.tables,
            start=1,
        ):
            rows: list[list[str]] = []

            for row in table.rows:
                cells = [
                    cell.text.strip().replace("\n", " ")
                    for cell in row.cells
                ]

                rows.append(cells)

            if not rows:
                continue

            markdown = self._table_to_markdown(rows)

            if markdown:
                pages.append(
                    ExtractedPage(
                        page_number=page_number,
                        content=markdown,
                        content_type="table",
                        metadata={
                            "source": "table",
                            "table_index": table_index,
                        },
                    )
                )

                page_number += 1

        return pages

    @staticmethod
    def _heading_level(style_name: str) -> int:
        """Extract the numeric heading level from a DOCX style name."""

        try:
            return int(style_name.split()[-1])
        except (ValueError, IndexError):
            return 1

    @staticmethod
    def _table_to_markdown(rows: list[list[str]]) -> str:
        """Convert DOCX table rows into Markdown."""

        if not rows:
            return ""

        column_count = max(len(row) for row in rows)

        normalized_rows = []

        for row in rows:
            normalized = list(row)
            normalized.extend([""] * (column_count - len(normalized)))

            normalized_rows.append(
                [
                    cell.replace("|", "\\|")
                    for cell in normalized
                ]
            )

        header = normalized_rows[0]
        separator = ["---"] * column_count

        lines = [
            "| " + " | ".join(header) + " |",
            "| " + " | ".join(separator) + " |",
        ]

        for row in normalized_rows[1:]:
            lines.append(
                "| " + " | ".join(row) + " |"
            )

        return "\n".join(lines)